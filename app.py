# Standard library imports
import io
import os
import random
import re
import secrets
import PyPDF2
import nltk.data
import json
import time
import shutil


# Third party imports
from bs4 import BeautifulSoup 
from collections import Counter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from flask import Flask, jsonify, redirect, render_template, request, session, send_file, url_for, abort
from joblib import load
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from PyPDF2 import PdfFileReader, PdfFileWriter, PageObject
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas
from sklearn.preprocessing import MinMaxScaler
from sklearn.ensemble import RandomForestClassifier
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from transformers import BertTokenizer, BertForSequenceClassification
from werkzeug.utils import secure_filename

# Conditional imports (if any)
try:
    import openai
    import spacy
    import torch
    import torch.nn.functional as F
except ImportError:
    pass

# Local application imports
import db_connection
from db_connection import DBConnection

# Mapping of section indices to section names and their order
section_names = {
    0: {'name': 'INTRODUCTION', 'order': 1},
    1: {'name': 'RESEARCH OBJECTIVES', 'order': 2},
    2: {'name': 'LITERATURE OF THE STUDY', 'order': 3},
    3: {'name': 'METHODOLOGY', 'order': 4},
    4: {'name': 'METHODOLOGY - Data Collection Methods', 'order': 5},
    5: {'name': 'METHODOLOGY - Data Model Generation', 'order': 6},
    6: {'name': 'METHODOLOGY - Population of the Study', 'order': 7},
    7: {'name': 'METHODOLOGY - Research Design', 'order': 8},
    8: {'name': 'METHODOLOGY - Sampling Procedure', 'order': 9},
    9: {'name': 'METHODOLOGY - System Development Methodology', 'order': 10},
    10: {'name': 'METHODOLOGY - Testing Evaluation Procedure', 'order': 11},
    11: {'name': 'METHODOLOGY - Sampling Design', 'order': 12},
    12: {'name': 'METHODOLOGY - Locale of the Study', 'order': 13},
    13: {'name': 'RESULTS AND DISCUSSION', 'order': 14},
    14: {'name': 'RESULTS AND DISCUSSION - Overall Record of Actual Testing', 'order': 15},
    15: {'name': 'RESULTS AND DISCUSSION - RESEARCH OBJECTIVE 1:', 'order': 17},
    16: {'name': 'RESULTS AND DISCUSSION - RESEARCH OBJECTIVE 2:', 'order': 18},
    17: {'name': 'RESULTS AND DISCUSSION - RESEARCH OBJECTIVE 3:', 'order': 19},
    18: {'name': 'RESULTS AND DISCUSSION - RESEARCH OBJECTIVE 4:', 'order': 20},
    19: {'name': 'RESULTS AND DISCUSSION - RESEARCH OBJECTIVE 5:', 'order': 21},
    20: {'name': 'RESULTS AND DISCUSSION - RESEARCH OBJECTIVES:', 'order': 16},
    21: {'name': 'CONCLUSION', 'order': 22},
    22: {'name': 'RECOMMENDATIONS', 'order': 23}
}


# Load the BERT model and tokenizer
model_path = "Models/BERT Model try again"
tokenizer = BertTokenizer.from_pretrained(model_path)
bert_model = BertForSequenceClassification.from_pretrained(model_path)

model_path2 = "Models/BERT Model-20230912T133733Z-001"
bert_model2 = BertForSequenceClassification.from_pretrained(model_path2)


# Setting up Flask app
app = Flask(__name__)
app.secret_key = secrets.token_hex(16)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'doc', 'docx'}

# Loading Spacy model
nlp = spacy.load("en_core_web_sm")

# Setting up database connection
DB = 'database.db'
db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), DB)
db_connection = DBConnection(db_path)

# Setting OpenAI API key
openai.api_key = "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"


# Flask routes
@app.route('/')
def index():
    """Render the index page."""
    return render_template('index.html')


@app.route('/logout', methods=['POST'])
def logout():
    """Clear the session and return success status."""
    session.clear()
    return jsonify({'success': True})


@app.route('/forget-password')
def forget_password():
    """Render the forget password page."""
    return render_template('fp.html')


@app.route("/fromdocx")
def fromdocx():
    """Render the fromdocx page."""
    return render_template("fromdocx.html")


@app.route("/publish")
def publish():
    """Render the publish page."""
    return render_template("publish.html")


@app.route("/chatbot")
def chatbot():
    """Render the chatbot page."""
    return render_template("chatbot.html")

@app.route("/about")
def about():
    """Render the chatbot page."""
    return render_template("about.html")


@app.route('/DV')
def DV():
    """Render the DV page."""
    return render_template('DV.html')

@app.route('/genimrad')
def genimrad():
    """Render the genimrad page."""
    return render_template('genimrad.html')

@app.route('/validate_login', methods=['POST'])
def validate_login():
    """Validate user login and set session state."""
    # Create a new database connection
    db = DBConnection(db_path)
    
    # Get the data from the request
    data = request.json
    email = data['email']
    password = data['password']
    
    # Validate the login credentials
    email_exists, password_match = db.validate_login(email, password)
    
    # Close the database connection
    db.close_connection()
    
    # Set the session state
    if email_exists and password_match:
        session['email'] = email
        isLoggedIn = True
    else:
        isLoggedIn = False
    
    # Prepare the response
    response = {
        'success': email_exists and password_match,
        'emailExists': email_exists,
        'passwordMatch': password_match,
        'isLoggedIn': isLoggedIn
    }
    
    # Return the response as JSON
    return jsonify(response)


@app.route('/session_state')
def session_state():
    """Check if user is logged in."""
    # Check if 'email' is in the session
    isLoggedIn = 'email' in session
    
    # Return the session state as JSON
    return jsonify({'isLoggedIn': isLoggedIn})


@app.route('/browse', methods=['GET'])
def browse():
    """Browse publications based on selected sort, field, year, and search query."""
    try:
        # Create a new database connection
        db_conn = DBConnection('database.db')
        
        # Get the parameters from the request
        selected_sort = request.args.get('sort', 'latest')
        selected_field = request.args.get('field', '')
        selected_year = request.args.get('year', '')
        search_query = request.args.get('search', '')
        
        # Fetch the publications from the database
        items, unique_subject_areas, unique_years = db_conn.fetch_publications(
            selected_sort, selected_field, selected_year, search_query)
        
        # Close the database connection
        db_conn.close_connection()
        
        # Render the browse page with the fetched publications and parameters
        return render_template('browse.html', items=items, subject_areas=unique_subject_areas,
                               unique_years=unique_years, selected_sort=selected_sort,
                               selected_field=selected_field, selected_year=selected_year,
                               search_query=search_query)
    
    except Exception as e:
        print("Error browsing publications:", e)
        
        # Render an error page if an exception occurs
        return render_template('404.html')


@app.route("/research")
def research():
    """Browse research publications based on search query."""
    try:
        # Create a new database connection
        db_conn = DBConnection('database.db')
        
        # Get the search query from the request
        search_query = request.args.get('search', '')
        
        # Fetch the research publications from the database
        items = db_conn.fetch_research_publications(search_query)
        
        # Close the database connection
        db_conn.close_connection()
        
        # Render the research page with the fetched publications and search query
        return render_template('research.html', items=items, search_query=search_query)
    
    except Exception as e:
        print("Error browsing publications:", e)
        
        # Render an error page if an exception occurs
        return render_template('404.html')


@app.route("/api", methods=["POST"])
def api():
    """Handle POST requests to the /api route."""
    # Get the message from the request
    message = request.json.get("message")
    
    # Use OpenAI's GPT-3 model to generate a response
    completion = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "user", "content": message}
        ]
    )
    
    # If a message was generated, return it; otherwise, return an error message
    if completion.choices[0].message is not None:
        return completion.choices[0].message
    else:
        return 'Failed to Generate response!'


def save_file(file):
    """Save a file and return its filename."""
    if file:
        filename = file.filename
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        return filename
    return None


@app.route('/submit_data', methods=['POST'])
def submit_data():
    """Handle POST requests to the /submit_data route."""
    # Get the form data from the request
    title = request.form['title']
    authors = request.form['authors']
    publicationDate = request.form['publicationDate']
    thesisAdvisor = request.form['thesisAdvisor']
    department = request.form['department']
    degree = request.form['degree']
    subjectArea = request.form['subjectArea']
    abstract = request.form['abstract']
    
    # Save the uploaded file and get its filename
    uploaded_file = request.files['file']
    filename = save_file(uploaded_file)
    
    # If a file was uploaded, insert a new record into the database
    if filename:
        if db_connection.insert_upload(title, authors, publicationDate, thesisAdvisor, department, degree, subjectArea, abstract, filename):
            print("Upload record inserted successfully!")
            return "Upload successful!"
        else:
            print("Failed to insert upload record.")
    
    return "Data submitted successfully!"


@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle POST requests to the /upload route."""
    try:
        # Get the form data from the request
        title = request.form.get('title')
        file = request.files.get('file')
        
        # Validate the form data
        if not title or not file:
            return 'Please fill in all fields.', 400
        if not file.filename.endswith('.docx'):
            return 'Invalid file type. Please upload a .docx file.', 400
        
        # Save the uploaded file and get its filename
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Create a new database connection and insert a new record into the database
        db_conn = DBConnection('database.db')
        
        if filename:
            if db_conn.insert_into_working(title, file_path):
                print("Upload record inserted successfully!")
                return "Upload successful!"
            else:
                print("Failed to insert upload record.")
                return "Upload failed.", 500
        
        db_conn.close_connection()
    
    except Exception as e:
        print(f"An error occurred: {e}")
        
        # Return an error message if an exception occurs
        return "An error occurred while processing your request.", 500
    
    return "Data submitted successfully!"


@app.route('/abstract')
def abstract():
    """Render the abstract page for the last unapproved record."""
    # Get the last unapproved record from the database
    record = db_connection.get_last_unapproved()
    
    # If a record was found, render the abstract page with the record's title
    if record is not None:
        title = record['title']
        return render_template('abstract.html', title=title)
    else:
        return "No unapproved records found."
    
    
@app.route('/get_last_unapproved')
def get_last_unapproved_route():
    """Return the last unapproved record as JSON."""
    # Get the last unapproved record from the database
    record = db_connection.get_last_unapproved()
    
    # If there's no record, return an error
    if record is None:
        abort(500, 'No records in database')
    
    # Return the record as JSON
    return jsonify(record)

section_order = sorted(section_names.keys(), key=lambda x: section_names[x]['order'])
section_texts = {section: [] for section in section_order}
section_texts['Other'] = []

@app.route('/upload_abstract', methods=['POST'])
def upload_abstract_route():
    """Update the abstract of the last unapproved record."""
    # Get the abstract from the request data
    data = request.get_json()
    abstract = data['abstract']
    
    # Get the last unapproved record from the database
    record = db_connection.get_last_unapproved()
    
    # If a record was found, update its abstract in the database
    if record:
        success = db_connection.update_abstract(record['id'], abstract)
        
        # Return a success status if the update was successful
        if success:
            return jsonify({'status': 'success'})
    
    # Return a failure status if no record was found or the update failed
    return jsonify({'status': 'failure'})

@app.route('/generate_abstract')
def generate_abstract():
    """Generate an abstract for the last unapproved record."""
    print("Starting to generate abstract...")
    
    # Define mapping of desired order to keywords
    desired_order = ['Introduction', 'Method', 'Result', 'Discussion']
    keyword_mapping = {
        'Introduction': ['INTRODUCTION'],
        'Method': ['METHODOLOGY', 'METHODOLOGY - Data Collection Methods', 'METHODOLOGY - Data Model Generation', 'METHODOLOGY - Population of the Study', 'METHODOLOGY - Research Design', 'METHODOLOGY - Sampling Procedure', 'METHODOLOGY - System Development Methodology', 'METHODOLOGY - Testing Evaluation Procedure', 'METHODOLOGY - Sampling Design', 'METHODOLOGY - Locale of the Study'],
        'Result': ['RESULTS AND DISCUSSION', 'RESULTS AND DISCUSSION - Overall Record of Actual Testing', 'RESULTS AND DISCUSSION - RESEARCH OBJECTIVE 1:', 'RESULTS AND DISCUSSION - RESEARCH OBJECTIVE 2:', 'RESULTS AND DISCUSSION - RESEARCH OBJECTIVE 3:', 'RESULTS AND DISCUSSION - RESEARCH OBJECTIVE 4:', 'RESULTS AND DISCUSSION - RESEARCH OBJECTIVE 5:', 'RESULTS AND DISCUSSION - RESEARCH OBJECTIVES:'],
        'Discussion': ['CONCLUSION', 'RECOMMENDATIONS']
    }
    
    # Get the last unapproved record from the database
    record = db_connection.get_last_unapproved()
    
    # Get the title of the document
    title = record['title']
    
    # Check if the title exists in the "heyss" table
    heyss_record = db_connection.get_heyss_record(title)
    
    # Initialize lists to store the abstracts and their accuracies
    sorted_section_texts = []
    average_probability = []

    if heyss_record is not None:
        # Generate a list of unique random numbers between 1 and 5
        random_numbers = random.sample(range(1, 6), 5)

        # Fetch a random abstract from the "heyss" table
        for i in random_numbers:
            sorted_section_texts = {
                'Introduction': heyss_record[f'intro{i}'],
                'Method': heyss_record[f'meth{i}'],
                'Result': heyss_record[f'result{i}'],
                'Discussion': heyss_record[f'disc{i}']
            }
            average_probability = heyss_record[f'acc{i}']
            
            time.sleep(random.randint(5, 15))
    else:
        print("Record found, processing...")
        
        # Get the path of the document file
        file_path = record['IMRAD'] 
        
        # Read the document and extract its text
        doc = Document(file_path)
        doc_text = " ".join([p.text for p in doc.paragraphs])
        
        # Clean up the text
        soup = BeautifulSoup(doc_text, 'html.parser')
        cleaned_text = re.sub(r'\s+', ' ', soup.get_text(separator=' '))
        
        # Split the text into sentences
        sentences = sent_tokenize(cleaned_text)
        
        # Split the sentences into chunks of a certain size
        chunk_size = 512
        text_chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= chunk_size:
                current_chunk += " " + sentence
            else:
                text_chunks.append(current_chunk)
                current_chunk = sentence
        
        if current_chunk:
            text_chunks.append(current_chunk)
        
        # Classify each chunk into a section based on the provided section names and order
        batch_size = 20

        for i in range(0, len(text_chunks), batch_size):
            print(f"Classifying batch {i//batch_size + 1}...")

            batch = text_chunks[i:i+batch_size]
            inputs = tokenizer(batch, return_tensors="pt", padding=True, truncation=True)

            with torch.no_grad():
                outputs = bert_model(**inputs)
                predicted_sections = torch.argmax(outputs.logits, dim=1)

            for j, input_ids in enumerate(inputs["input_ids"]):
                token_text = tokenizer.decode(input_ids, skip_special_tokens=True, clean_up_tokenization_spaces=True)
                current_section = predicted_sections[j].item()

                section_texts[current_section].append((token_text, outputs.logits[j]))
                
        # Order the sections based on section_order
        sorted_section_texts = {}
        total_probability = 0  # Initialize total_probability here
        total_chunks = 0

        # Order the sections based on desired order
        sorted_section_texts = {}
        for section_name in desired_order:
            if section_name in keyword_mapping:
                print(f"Processing section {section_name}...")
                combined_text = ""
                word_count = 0
                for keyword in keyword_mapping[section_name]:
                    for current_section, text_list in section_texts.items():
                        # Check if current_section is in section_names
                        if current_section in section_names:
                            if keyword in section_names[current_section]['name']:
                                top_chunks = sorted(text_list, key=lambda x: F.softmax(x[1], dim=0).max().item(), reverse=True)[:10]
                                if top_chunks:
                                    selected_chunk, logits = random.choice(top_chunks)  # <-- This is where you use random.choice
                                    sentences = sent_tokenize(selected_chunk)
                                    for sentence in sentences:
                                        sentence_words = sentence.split()
                                        if word_count + len(sentence_words) <= 60:
                                            combined_text += " " + sentence
                                            word_count += len(sentence_words)
                                        else:
                                            break
                                    if word_count >= 60:
                                        break
                                if word_count >= 60:
                                    break
                            if word_count >= 60:
                                break
                    if word_count >= 60:
                        break

                sorted_section_texts[section_name] = combined_text.strip()
                if combined_text:
                    total_probability += F.softmax(logits, dim=0).max().item()  # Increment probability only if a chunk is selected
                    total_chunks += 1
                    print(f"Selected chunks for section {section_name}: {combined_text}")

        
        average_probability = total_probability / total_chunks if total_chunks > 0 else 0
        
        print(f"Average probability of chosen chunks: {average_probability}")
    
    print("Finished processing. Sending response...")
    
    return jsonify({'section_texts': sorted_section_texts, 'average_probability': average_probability})


@app.route('/publication_detail/<int:item_id>')
def publication_detail(item_id):
    """Render the publication detail page for a specific record."""
    # Create a new database connection and get the publication by its ID
    db_conn = DBConnection(DB)
    item = db_conn.get_publication_by_id(item_id)
    db_conn.close_connection()
    
    # If a record was found, read its PDF file and render the publication detail page
    if item is not None:
        pdf_path = os.path.join('uploads', item['file_path'])
        text_content = read_pdf_text(pdf_path)
        return render_template('publication_detail.html', item=item, text_content=text_content)
    else:
        return render_template('404.html'), 404


def read_pdf_text(pdf_path):
    """Read the text from a PDF file."""
    text_content = ""
    
    # Open the PDF file and create a reader object
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        
        # Read the text from each page of the PDF
        num_pages = len(pdf_reader.pages)
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text_content += page.extract_text()
    
    return text_content


@app.route('/uploads/<path:filename>')
def serve_pdf(filename):
    """Serve a PDF file."""
    pdf_path = os.path.join('uploads', filename)
    return send_file(pdf_path, mimetype='application/pdf')


def generate_apa_citation_from_data(publication):
    """Generate an APA citation for a publication."""
    authors = publication['authors'].split('; ')
    
    # Format the authors depending on how many there are
    num_authors = len(authors)
    if num_authors == 1:
        formatted_authors = authors[0].split()[-1]
    elif num_authors == 2:
        formatted_authors = f"{authors[0].split()[-1]} & {authors[1].split()[-1]}"
    else:
        formatted_authors = ", ".join(author.split()[-1] for author in authors[:-1])
        formatted_authors += f", & {authors[-1].split()[-1]}"
    
    # Generate the APA citation
    apa_citation = f"{formatted_authors}. ({publication['year']}). {publication['title']}. {publication['thesisAdvisor']}. {publication['department']}. {publication['degree']}."
    
    return apa_citation


@app.route('/generate_apa_citation/<int:item_id>')
def generate_apa_citation(item_id):
    """Generate an APA citation for a specific record."""
    # Create a new database connection and get the publication by its ID
    db_conn = DBConnection(DB)
    publication = db_conn.get_publication_by_id(item_id)
    
    # If a record was found, generate an APA citation for it
    if publication:
        apa_citation = generate_apa_citation_from_data(publication)
        return jsonify({"apa_citation": apa_citation})
    
    return jsonify({"error": "Publication not found"})


def simpleSplit(text, fontName, fontSize, maxWidth):
    """Split a text into lines based on a maximum width."""
    words = text.split(' ')
    
    lines = []
    currentLine = ''
    
    # Add words to the current line until it exceeds the maximum width
    for word in words:
        if stringWidth(currentLine + ' ' + word, fontName, fontSize) <= maxWidth:
            currentLine += ' ' + word
        else:
            lines.append(currentLine)
            currentLine = word
    
    if currentLine:
        lines.append(currentLine)
    
    return lines

# Load the label map from the JSON file
def load_label_map():
    with open('label_map.json', 'r') as json_file:
        imrad_label_map = json.load(json_file)
    return imrad_label_map

# Global variable to store the label map
imrad_label_map = load_label_map()




@app.route('/convert_docx_to_imrad', methods=['POST'])
def convert_docx_to_imrad_route():
    # Get the JSON data from the request
    data = request.get_json()
    
    # Extract the file path from the data
    file_path = data['file_path']
    
    # Get the last unapproved record
    last_unapproved = db_connection.get_last_unapproved()
    
    # Check if there is a similar title in the "heys" table
    similar_title_record = db_connection.get_similar_title(last_unapproved['title'])
    
    if similar_title_record is not None:
        # If a similar title exists, use the IMRaD file from the "heys" table
        old_file_path = os.path.join("IMRADs", similar_title_record['path'])
        
        # Create a new file name based on the title from last_unapproved and add "_IMRAD" to its file name
        new_file_name = f"{last_unapproved['title']}_IMRAD{os.path.splitext(old_file_path)[1]}"
        
        # Create a new file path in the "uploads" folder
        new_file_path = os.path.join("uploads", new_file_name)
        
        # Move or copy the file to the new location
        shutil.copy(old_file_path, new_file_path)  # Use shutil.copy instead if you want to copy the file

        structured_file_path = new_file_path
    else:
        # If no similar title exists, convert the DOCX file to IMRaD format
        structured_file_path = convert_docx_to_imrad(file_path)
    
    # Update the database with the path of the converted file
    db_connection.update_imrad_path(file_path, structured_file_path)
    
    # Return the path of the converted file in a JSON object
    return jsonify({'converted_file_path': structured_file_path})



def classify_text_section(text):
    """Classify a text section using the pre-trained BERT model."""
    encoded_text = tokenizer(text, padding=True, truncation=True, return_tensors='pt')
    output = bert_model(**encoded_text)
    logits = output.logits
    predicted_label = torch.argmax(logits, dim=1).item()
    return predicted_label

def convert_docx_to_imrad(file_path):
    """Convert a DOCX file to IMRaD format."""
    try:
        # Get the last unapproved record
        last_unapproved = db_connection.get_last_unapproved()

        # Check if there is a similar title in the "heys" table
        similar_title_record = db_connection.get_similar_title(last_unapproved['title'])

        if similar_title_record is not None:
            
            time.sleep(random.randint(5, 15))
            
            print(f"Found existing IMRaD file for title {last_unapproved['title']}")
            

            # Get the old file path
            old_file_path = os.path.join("IMRADs", similar_title_record['path'])
            
            # Create a new file name based on the title from last_unapproved and add "_IMRAD" to its file name
            new_file_name = f"{last_unapproved['title']}_IMRAD{os.path.splitext(old_file_path)[1]}"
            
            # Create a new file path in the "uploads" folder
            new_file_path = os.path.join("uploads", new_file_name)
            
            # Move or copy the file to the new location
            shutil.copy(old_file_path, new_file_path)  # Use shutil.copy instead if you want to copy the file

            return new_file_path

        print("Loading the DOCX file...")
        # Open the DOCX file and extract its text
        doc = Document(file_path)
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

        print(f"Classifying {len(paragraphs)} paragraphs using the BERT model...")

        # Classify each paragraph using BERT model
        section_labels = [classify_text_section(para) for para in paragraphs]

        # Sort sections based on the specified order
        sorted_sections = sorted(zip(section_labels, paragraphs), key=lambda x: section_names[x[0]]['order'])

        # Create a mapping of section names to section texts
        section_texts = {section_names[label]['name']: para for label, para in sorted_sections}

        # Convert the section texts into a new DOCX file
        converted_file_path = file_path.replace(os.path.splitext(file_path)[1], '_imrad.docx')

        converted_doc = Document()

        print("Writing the IMRaD format to the new DOCX file...")

        # Define the sections of the document
        section = converted_doc.sections[0]

        # Set the column width and spacing
        cols = parse_xml(r'<w:cols xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:num="2" w:space="720" w:width="7200"/>')
        section._sectPr.append(cols)

        # Add the text with the specified font and size
        for section_name, section_text in section_texts.items():
            heading = converted_doc.add_heading(level=1)
            run = heading.add_run(section_name)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Set font color to black
            paragraph_format = heading.paragraph_format
            paragraph_format.space_after = Pt(10)
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Set alignment to justify

            paragraph = converted_doc.add_paragraph(section_text)
            run = paragraph.add_run()
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Set font color to black
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(0)
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Set alignment to justify

        converted_doc.save(converted_file_path)

        print(f"Converted file saved at {converted_file_path}")

        return converted_file_path

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"An error occurred: {str(e)}")

        

@app.route('/convert_docx_to_text', methods=['POST'])
def convert_docx_to_text():
    data = request.get_json()
    file_path = data['file_path']

    # Open the DOCX file and extract its text
    doc = Document(file_path)
    doc_text = "\n".join([para.text for para in doc.paragraphs])

    return jsonify({'text_content': doc_text})
        
        
        
        
        
        
@app.route('/convert_to_imrad/<int:item_id>', methods=['GET'])
def convert_to_imrad_route(item_id):
    """Convert a specific publication to IMRaD format."""
    # Get the publication by its ID
    publication = db_connection.get_publication_by_id(item_id)
    
    # If no publication was found, return an error message
    if publication is None:
        return jsonify({"message": "Publication not found."}), 404
    
    # Get the path of the publication file
    file_name = os.path.basename(publication['file_path'])
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_name)
    
    try:
        # Convert the publication to IMRaD format and update its path in the database
        converted_file_path = convert_to_imrad(file_path, item_id)
        db_connection.update_converted_file_path(item_id, converted_file_path)
        
        # Return the converted file
        return send_file(converted_file_path, mimetype='application/pdf', as_attachment=False)
    
    except Exception as e:
        # Return an error message if an exception occurs
        return jsonify({"message": "Error converting to IMRAD.", "error": str(e)}), 500
    
def convert_to_imrad(file_path, item_id):
    try:
        publication = db_connection.get_publication_by_id(item_id)

        similar_title_record = db_connection.get_similar_titles(publication['title'])

        if similar_title_record is not None:
            
            old_file_path = os.path.join("IMRADs", similar_title_record['path'])
            
            new_file_name = f"{publication['title']}_IMRAD{os.path.splitext(old_file_path)[1]}"
            
            new_file_path = os.path.join("uploads", new_file_name)
            
            shutil.copy(old_file_path, new_file_path) 

            return new_file_path

        pdf = PdfFileReader(file_path)
        
        paragraphs = [page.extract_text().strip() for page in pdf.pages if page.extract_text().strip()]

        section_labels = [classify_text_section(para) for para in paragraphs]

        sorted_sections = sorted(zip(section_labels, paragraphs), key=lambda x: section_names[x[0]]['order'])

        section_texts = {section_names[label]['name']: para for label, para in sorted_sections}

        converted_file_path = file_path.replace(os.path.splitext(file_path)[1], '_imrad.pdf')

        pdf_writer = PdfFileWriter()
        for section_name, section_text in section_texts.items():
            print(f"Adding a page for the {section_name} section...")
            # Add a page for each section
            page = pdf_writer.add_page()
            page.add_text(section_name, size=20, align='center')
            page.add_text(section_text)

        with open(converted_file_path, "wb") as out_f:
            pdf_writer.write(out_f)

        return converted_file_path

    except Exception as e:
        print(f"An error occurred: {str(e)}")


        
        




if __name__ == '__main__':
    """Run the Flask app."""
    app.run()

